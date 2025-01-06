from langchain.schema import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import os
import re
from typing import List, Optional, Dict, Any, Union
from utils.logging_config import logger
from utils.error_handler import DocumentProcessingError, handle_error
from dspy_pipeline import IEPPipeline

def load_pdf(file_path: str) -> List[Document]:
    """Load and process a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        List of Document objects containing the PDF content
        
    Raises:
        DocumentProcessingError: If the PDF cannot be processed
    """
    try:
        reader = PdfReader(file_path)
        documents = []
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(file_path),
                        "page": i + 1,
                        "type": "pdf"
                    }
                ))
        return documents
    except Exception as e:
        raise DocumentProcessingError(
            f"Failed to process PDF file: {file_path}",
            details={"error": str(e)}
        )

def load_docx(file_path: str) -> List[Document]:
    """Load and process a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        List of Document objects containing the DOCX content
        
    Raises:
        DocumentProcessingError: If the DOCX cannot be processed
    """
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        if text.strip():
            return [Document(
                page_content=text,
                metadata={
                    "source": os.path.basename(file_path),
                    "type": "docx"
                }
            )]
        return []
    except Exception as e:
        raise DocumentProcessingError(
            f"Failed to process DOCX file: {file_path}",
            details={"error": str(e)}
        )

def load_documents(file_paths: List[str]) -> List[Document]:
    """Load documents from multiple file paths.
    
    Args:
        file_paths: List of paths to documents
        
    Returns:
        List of processed Document objects
    """
    documents: List[Document] = []
    
    for file_path in file_paths:
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                documents.extend(load_pdf(file_path))
            elif ext in ['.docx', '.doc']:
                documents.extend(load_docx(file_path))
            else:
                logger.warning(f"Unsupported file type: {ext}")
                
        except DocumentProcessingError as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            handle_error(e, {"file_path": file_path})
        except Exception as e:
            logger.error(f"Unexpected error processing file {file_path}: {str(e)}")
            handle_error(e, {"file_path": file_path})
            
    return documents

def process_with_dspy(documents: List[Document]) -> List[Document]:
    """Process documents using the DSPy pipeline.
    
    Args:
        documents: List of documents to process
        
    Returns:
        List of processed Document objects
    """
    iep_pipeline = IEPPipeline()
    results: List[Document] = []

    for doc in documents:
        try:
            processed_docs = iep_pipeline.process_documents([doc])
            
            if processed_docs and len(processed_docs) > 0:
                iep_result = processed_docs[0]
                results.append(iep_result)
                logger.debug(f"Successfully processed document: {doc.metadata.get('source')}")
            else:
                logger.warning(f"No processing results for document: {doc.metadata.get('source')}")
                results.append(doc)
            
        except Exception as e:
            error_info = handle_error(e, {
                "document_source": doc.metadata.get("source"),
                "document_type": doc.metadata.get("type")
            })
            logger.error(f"Error processing document with DSPy: {error_info['message']}")
            results.append(doc)

    return results