import unittest
from unittest.mock import Mock, patch
from langchain.schema import Document
from utils.error_handler import DocumentProcessingError
from loaders import load_documents, process_with_dspy
import os
import tempfile

class TestDocumentProcessing(unittest.TestCase):
    def setUp(self):
        """Set up test environment."""
        self.test_dir = tempfile.mkdtemp()
        
    def tearDown(self):
        """Clean up test environment."""
        import shutil
        shutil.rmtree(self.test_dir)
        
    def create_test_pdf(self, content: str) -> str:
        """Create a test PDF file."""
        from reportlab.pdfgen import canvas
        pdf_path = os.path.join(self.test_dir, "test.pdf")
        c = canvas.Canvas(pdf_path)
        c.drawString(72, 720, content)
        c.save()
        return pdf_path
        
    def create_test_docx(self, content: str) -> str:
        """Create a test DOCX file."""
        from docx import Document
        docx_path = os.path.join(self.test_dir, "test.docx")
        doc = Document()
        doc.add_paragraph(content)
        doc.save(docx_path)
        return docx_path

    def test_load_pdf(self):
        """Test PDF loading functionality."""
        test_content = "Test PDF Content"
        pdf_path = self.create_test_pdf(test_content)
        
        documents = load_documents([pdf_path])
        self.assertTrue(len(documents) > 0)
        self.assertIn("pdf", documents[0].metadata["type"])
        
    def test_load_docx(self):
        """Test DOCX loading functionality."""
        test_content = "Test DOCX Content"
        docx_path = self.create_test_docx(test_content)
        
        documents = load_documents([docx_path])
        self.assertEqual(len(documents), 1)
        self.assertIn("docx", documents[0].metadata["type"])
        
    def test_invalid_file(self):
        """Test handling of invalid files."""
        documents = load_documents(["nonexistent.pdf"])
        self.assertEqual(len(documents), 0)
        
    def test_empty_file_list(self):
        """Test handling of empty file list."""
        documents = load_documents([])
        self.assertEqual(len(documents), 0)
        
    @patch('loaders.IEPPipeline')
    def test_dspy_processing(self, mock_pipeline):
        """Test DSPy processing functionality."""
        # Mock the pipeline
        mock_instance = Mock()
        mock_instance.process_documents.return_value = [
            Document(
                page_content="Processed content",
                metadata={"source": "test.pdf", "processed": True}
            )
        ]
        mock_pipeline.return_value = mock_instance
        
        # Test document
        test_doc = Document(
            page_content="Test content",
            metadata={"source": "test.pdf"}
        )
        
        # Process document
        results = process_with_dspy([test_doc])
        
        # Verify results
        self.assertEqual(len(results), 1)
        self.assertTrue(results[0].metadata.get("processed"))
        
    def test_error_handling(self):
        """Test error handling during document processing."""
        # Create an invalid PDF file
        with open(os.path.join(self.test_dir, "invalid.pdf"), "w") as f:
            f.write("Not a PDF file")
            
        # Attempt to load the invalid file
        documents = load_documents([os.path.join(self.test_dir, "invalid.pdf")])
        self.assertEqual(len(documents), 0)

if __name__ == '__main__':
    unittest.main() 